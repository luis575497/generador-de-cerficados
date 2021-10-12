# Documentacion de la librería docx - https://python-docx.readthedocs.io/
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
from docx.shared import Pt
from os import remove

def crear_certificado(estudiante,cedula,facultad,carrera,handle,especialista):
    # Funcion para poner el mes
    def mes(num):
        meses = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']
        return meses[num-1]

    # Datos
    fecha_actual = datetime.datetime.now()
    fecha = f'{fecha_actual.day} de {mes(fecha_actual.month)} de {fecha_actual.year}'

    # Creación de la instancia de clase Document
    document = Document()

    # Encabezado
    document.add_picture('recursos/ucuenca.jpg', width=Inches(3))
    last_paragraph = document.paragraphs[0]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    encabezado = document.add_paragraph()
    texto_encabezado = encabezado.add_run("\n\n\nC E R T I F I CA\n\n\n")
    texto_encabezado.bold = True
    fuente_texto_encabezado = texto_encabezado.font
    fuente_texto_encabezado.size = Pt(12)
    fuente_texto_encabezado.name = 'Arial'
    formato_encabezado = encabezado.paragraph_format
    formato_encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Texto central
    parrafo = document.add_paragraph()
    parrafo1 = parrafo.add_run('Que ')
    parrafo2 = parrafo.add_run(f'{estudiante.upper()} ')
    parrafo3 = parrafo.add_run(' con cédula No. ')
    parrafo4 = parrafo.add_run(f'{cedula}, ')
    parrafo5 = parrafo.add_run(f' estudiante de la {facultad}, {carrera}, no adeuda ningún bien, ni material bibliográfico en esta dependencia.')
    parrafo.add_run('\n')
    parrafo.add_run('\n')
    parrafo.add_run('\n')
    formato_parrafo = parrafo.paragraph_format
    formato_parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fuente_parrafo1 = parrafo1.font
    fuente_parrafo1.size = Pt(12)
    fuente_parrafo1.name = 'Arial'
    parrafo2.bold = True
    fuente_parrafo2 = parrafo2.font
    fuente_parrafo2.size = Pt(12)
    fuente_parrafo2.name = 'Arial'
    fuente_parrafo3 = parrafo3.font
    fuente_parrafo3.size = Pt(12)
    fuente_parrafo3.name = 'Arial'
    parrafo4.bold = True
    fuente_parrafo4 = parrafo4.font
    fuente_parrafo4.size = Pt(12)
    fuente_parrafo4.name = 'Arial'
    fuente_parrafo5 = parrafo5.font
    fuente_parrafo5.size = Pt(12)
    fuente_parrafo5.name = 'Arial'

    # Ciudad y Fecha
    ciudad_fechas = document.add_paragraph()
    ciudadyfecha = ciudad_fechas.add_run(f'Cuenca, {fecha}')
    fuente_ciudad_fecha = ciudadyfecha.font
    fuente_ciudad_fecha.size = Pt(12)
    fuente_ciudad_fecha.name = "Arial"
    formato_ciudad_fechas = ciudad_fechas.paragraph_format
    formato_ciudad_fechas.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ciudad_fecha = document.add_paragraph()
    ciudad_fecha.add_run('\n')
    ciudad_fecha.add_run('\n')
    atentamente = ciudad_fecha.add_run('Atentamente,')
    ciudad_fecha.add_run('\n')
    ciudad_fecha.add_run('\n')
    ciudad_fecha.add_run('\n')
    ciudad_fecha.add_run('\n')
    ciudad_fecha.add_run('\n')
    ciudad_fecha.add_run('\n')
    font_atentamente = atentamente.font
    font_atentamente.size = Pt(12)
    font_atentamente.name = "Arial"
    formato_ciudad_fecha = ciudad_fecha.paragraph_format
    formato_ciudad_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Firma del especialista
    firma = document.add_paragraph()
    firma_especialista = firma.add_run(f"{especialista}\n\n")
    font_especialista = firma_especialista.font
    font_especialista.size = Pt(12)
    font_especialista.name = "Arial"
    formato_firma = firma.paragraph_format
    formato_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parte final del documento
    link_tt = document.add_paragraph()
    link = link_tt.add_run(f'{handle}')
    font_link = link.font
    font_link.size = Pt(12)
    font_link.name = "Arial"
    font_link.bold = True
    nota = link_tt.add_run('\n\nNota: Certificado digital emitido de forma extraordinaria por la emergencia sanitaria COVID-19; razón por la cual no consta el sello de la dependencia.')
    font_nota = nota.font
    font_nota.name = "Arial"
    font_nota.size = Pt(12)

    # Crear el documento si no existe previamente
    try:
        document.save(f'Certif. No Adeud - {estudiante}.docx')

    except:
        remove(f'Certif. No Adeud - {estudiante}.docx')
        document.save(f'Certif. No Adeud - {estudiante}.docx')


